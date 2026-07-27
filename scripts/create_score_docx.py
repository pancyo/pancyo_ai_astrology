from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'docs' / '点数計算方法.md'
OUTPUT = ROOT / 'docs' / '点数計算方法.docx'

FONT = 'Meiryo'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = '1F2937'
HEADER_FILL = 'E8EEF5'
GRID = 'B8C4D3'
WIDTH_DXA = 9360


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:ascii'), FONT)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width))
    tc_w.set(qn('w:type'), 'dxa')


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement('w:cantSplit')
    tr_pr.append(marker)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement('w:tblHeader')
    marker.set(qn('w:val'), 'true')
    tr_pr.append(marker)


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in('w:tblLayout')
    if tbl_layout is None:
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    indent = OxmlElement('w:tblInd')
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in('w:tcMar')
            if margins is None:
                margins = OxmlElement('w:tcMar')
                tc_pr.append(margins)
            for side in ('top', 'bottom', 'start', 'end'):
                node = margins.find(qn(f'w:{side}'))
                if node is None:
                    node = OxmlElement(f'w:{side}')
                    margins.append(node)
                node.set(qn('w:w'), '80' if side in ('top', 'bottom') else '120')
                node.set(qn('w:type'), 'dxa')


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, end])
    set_run_font(run, size=9, color='6B7280')


def add_inline(paragraph, text, size=11, color=INK):
    pieces = re.split(r'(`[^`]+`|\*\*[^*]+\*\*)', text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith('`') and piece.endswith('`'):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, size=size, color=DARK_BLUE)
        elif piece.startswith('**') and piece.endswith('**'):
            run = paragraph.add_run(piece[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(piece)
            set_run_font(run, size=size, color=color)


def set_paragraph(paragraph, before=0, after=6, line=1.25, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_table(doc, rows):
    columns = len(rows[0])
    table = doc.add_table(rows=0, cols=columns)
    table.style = 'Table Grid'
    if columns == 2:
        widths = [2700, 6660]
    elif columns == 3:
        widths = [2250, 1800, 5310]
    elif columns == 5:
        widths = [3000, 1200, 1200, 1700, 2260]
    else:
        widths = [1800, 1700, 1700, 4160]
    for row_index, source_row in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for index, source_cell in enumerate(source_row):
            cell = row.cells[index]
            cell.text = ''
            paragraph = cell.paragraphs[0]
            set_paragraph(paragraph, after=0, line=1.12)
            add_inline(paragraph, source_cell, size=8.7 if columns >= 3 else 9)
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        if row_index == 0:
            repeat_header(row)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines, index):
    raw_rows = []
    while index < len(lines) and lines[index].strip().startswith('|'):
        line = lines[index].strip()
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        if not all(re.fullmatch(r'[-: ]+', cell) for cell in cells):
            raw_rows.append(cells)
        index += 1
    return raw_rows, index


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 18, 10),
        ('Heading 2', 13, BLUE, 14, 7),
        ('Heading 3', 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(header, after=0)
    add_inline(header, 'ぱんちょ式星占い | 点数計算方法', size=9, color='6B7280')
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(footer, after=0)
    add_inline(footer, 'ぱんちょ式星占い  |  ', size=9, color='6B7280')
    add_page_field(footer)

    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    index = 0
    title_done = False
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith('|'):
            rows, index = parse_table(lines, index)
            if rows:
                add_table(doc, rows)
            continue
        if line.startswith('# '):
            if not title_done:
                paragraph = doc.add_paragraph()
                set_paragraph(paragraph, before=18, after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                add_inline(paragraph, line[2:], size=24, color=INK)
                for run in paragraph.runs:
                    run.bold = True
                title_done = True
            else:
                paragraph = doc.add_paragraph(style='Heading 1')
                add_inline(paragraph, line[2:], size=16, color=BLUE)
            index += 1
            continue
        if line.startswith('## '):
            paragraph = doc.add_paragraph(style='Heading 1')
            add_inline(paragraph, line[3:], size=16, color=BLUE)
            index += 1
            continue
        if line.startswith('### '):
            paragraph = doc.add_paragraph(style='Heading 2')
            add_inline(paragraph, line[4:], size=13, color=BLUE)
            index += 1
            continue
        if line.startswith('- '):
            paragraph = doc.add_paragraph(style='List Bullet')
            set_paragraph(paragraph, after=4, line=1.25)
            add_inline(paragraph, line[2:], size=11)
            index += 1
            continue
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, after=6, line=1.25)
        add_inline(paragraph, line, size=11)
        index += 1

    doc.core_properties.title = 'ぱんちょ式星占い 点数計算方法'
    doc.core_properties.subject = '毎日・週・月・年占いの点数計算仕様'
    doc.core_properties.author = 'ぱんちょ式星占い'
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    build()
